from flask_cors import CORS
from flask import Flask, jsonify, request, render_template, send_from_directory, session, send_file
import warnings
import csv
import os
import pandas as pd
from scipy.interpolate import lagrange
from sklearn.cluster import KMeans
from sklearn import preprocessing
from sklearn.preprocessing import MinMaxScaler
from docx import Document
from pdfminer.pdfparser import PDFParser, PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from sklearn.ensemble import RandomForestClassifier
from sklearn import tree
import random
from sklearn.svm import SVC
from sklearn.svm import SVR
from wordcloud import WordCloud, STOPWORDS, ImageColorGenerator
import numpy as np
from PIL import Image
from sklearn.linear_model import LinearRegression
from sklearn.linear_model import LogisticRegression
from sklearn.decomposition import PCA as sklearnPCA
from scipy import stats
from sklearn.manifold import TSNE as sklearnTSNE
import jieba.analyse
warnings.filterwarnings("ignore")
charset = "utf8"

app = Flask(__name__)
app.config['UPLOAD_FOLDER']=r"./upload"
CORS(app)

def ploy(s,n,k=3):
    y=s[list(range(n-k,n))+list(range(n+1,n+1+k))]#取数
    y=y[y.notnull()]
    return lagrange(y.index,list(y))(n)

def train_valid_data(data_):
    filesize = int(0.9 * len(data_))
    # 训练集和测试集的比例为9:1
    train_data_ = [each[0:-1] for each in data_[:filesize]]
    train_target_ = [each[-1] for each in data_[:filesize]]

    test_data_ = [each[0:-1] for each in data_[filesize:]]
    test_target_ = [each[-1] for each in data_[filesize:]]

    return train_data_, train_target_, test_data_, test_target_

def train_lr_valid_data(data_):
    filesize = int(0.9 * len(data_))
    # 训练集和测试集的比例为9:1
    train_data_ = [each[0:-1] for each in data_[:filesize]]
    train_target_ = [int(each[-1]) for each in data_[:filesize]]

    test_data_ = [each[0:-1] for each in data_[filesize:]]
    test_target_ = [int(each[-1]) for each in data_[filesize:]]

    return train_data_, train_target_, test_data_, test_target_


def get_train_dataset(filename):
    data = []
    with open('./upload/'+filename, 'r', newline='') as csvFile:
        csv_file = csv.reader(csvFile)
        for content in csv_file:
            content = list(map(float,content))
            if len(content)!=0:
                data.append(content)
    random.shuffle(data)
    return data

def get_test_dataset(filename):
    data = []
    with open('./upload/'+filename, 'r', newline='') as csvFile:
        csv_file = csv.reader(csvFile)
        for content in csv_file:
            content = list(map(float,content))
            if len(content)!=0:
                data.append(content)
    return data

@app.route('/kmeans/<argv>/<ran>', methods=['post', 'get'])
def kmeans(argv,ran):
    if argv == "download":
        path = r'./download/result.csv'
        rename = "result.csv"
        try:
            return send_file(path, mimetype='.csv', attachment_filename=rename, as_attachment=True)
        except Exception as E:
            return "error"
    #接受文件  部分
    clusters = int(argv)
    print('上传参数：',argv)
    file_obj = request.files['file']
    if file_obj is None:
        return "nofile"

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj.filename)
    file_obj.save(file_path)
    #处理数据
    data = []
    with open('./upload/'+file_obj.filename, 'r', newline='') as csvFile:
        csv_file = csv.reader(csvFile)
        for content in csv_file:
            content = list(map(float, content))
            if len(content) != 0:
                data.append(content)

    #聚类
    estimator = KMeans(n_clusters=clusters)  # 构造聚类器
    estimator.fit(data)  # 聚类
    label_pred = estimator.labels_  # 获取聚类标签
    centroids = estimator.cluster_centers_  # 获取聚类中心
    with open(r'./download/result.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        csv_writer.writerow(["label"])
        for line in label_pred:
            # print(line)
            temp = str(line)
            # print(temp)
            csv_writer.writerow([temp])
        csv_writer.writerow(["clusterCenter"])
        for line in centroids:
            # print(line)
            csv_writer.writerow(line)

    return 'success'


@app.route('/pca/<argv>/<ran>', methods=['post', 'get'])
def pca(argv,ran):
    if argv == "download":
        path = r'./download/result.csv'
        rename = "result.csv"
        try:
            return send_file(path, mimetype='.csv', attachment_filename=rename, as_attachment=True)
        except Exception as E:
            return "error"
    #接受文件  部分
    pca_dim = int(argv)
    print('维度：',argv)
    file_obj = request.files['file']
    if file_obj is None:
        return "nofile"

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj.filename)
    file_obj.save(file_path)
    #处理数据
    data = []
    traffic_feature = []
    traffic_target = []
    with open('./upload/'+file_obj.filename, 'r', newline='') as csvFile:
        csv_file = csv.reader(csvFile)
        for content in csv_file:
            content = list(map(float, content))
            if len(content) != 0:
                data.append(content)
                traffic_feature.append(content[0:-2])
                traffic_target.append(content[-1])
    #聚类

    min_max_scaler = preprocessing.MinMaxScaler()
    traffic_feature = min_max_scaler.fit_transform(traffic_feature)

    sklearn_pca = sklearnPCA(n_components=pca_dim)
    sklearn_transf = sklearn_pca.fit_transform(traffic_feature)

    with open(r'./download/result.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in sklearn_transf:
            csv_writer.writerow(line)
    return 'success'


@app.route('/TSNE/<argv>/<ran>', methods=['post', 'get'])
def TSNE(argv,ran):
    if argv == "download":
        path = r'./download/result.csv'
        rename = "result.csv"
        try:
            return send_file(path, mimetype='.csv', attachment_filename=rename, as_attachment=True)
        except Exception as E:
            return "error"
    #接受文件  部分
    TSNE_dim = int(argv)
    print('维度：',argv)
    file_obj = request.files['file']
    if file_obj is None:
        return "nofile"

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj.filename)
    file_obj.save(file_path)
    #处理数据
    data = []
    traffic_feature = []
    traffic_target = []
    with open('./upload/'+file_obj.filename, 'r', newline='') as csvFile:
        csv_file = csv.reader(csvFile)
        for content in csv_file:
            content = list(map(float, content))
            if len(content) != 0:
                data.append(content)
                traffic_feature.append(content[0:-2])
                traffic_target.append(content[-1])
    #聚类

    min_max_scaler = preprocessing.MinMaxScaler()
    traffic_feature = min_max_scaler.fit_transform(traffic_feature)

    sklearn_TSNE = sklearnTSNE(n_components=TSNE_dim)
    sklearn_transf = sklearn_TSNE.fit_transform(traffic_feature)

    with open(r'./download/result.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in sklearn_transf:
            csv_writer.writerow(line)
    return 'success'


@app.route('/Ttest/<argv>/<ran>', methods=['post', 'get'])
def Ttest(argv,ran):
    if argv == "download":
        path = r'./download/result.csv'
        rename = "result.csv"
        try:
            return send_file(path, mimetype='.csv', attachment_filename=rename, as_attachment=True)
        except Exception as E:
            return "error"
    #接受文件  部分
    Tnum = float(argv)
    file_obj = request.files['file']
    if file_obj is None:
        return "nofile"

    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj.filename)
    file_obj.save(file_path)
    #处理数据
    data = []
    with open('./upload/'+file_obj.filename, 'r', newline='') as csvFile:
        csv_file = csv.reader(csvFile)
        for content in csv_file:
            content = list(map(float, content))
            if len(content) != 0:
                data.append(content)
    #聚类
    result = stats.ttest_1samp(data, Tnum)
    print('argv:',argv)
    print('data:', data)
    with open(r'./download/result.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        csv_writer.writerow(result[1])
        if result[1] < 0.05:
            csv_writer.writerow(["The difference is significant"])
        else:
            csv_writer.writerow(["The difference is not significant"])
    return 'success'


@app.route('/pdf2word/<argv>/<ran>', methods=['post', 'get'])
def pdf2word(argv,ran):
    # pip install pdfminer3k
    # pip install python_docx
    if argv == "download":
        path = r'./download/result.docx'
        rename = "result.docx"
        try:
            return send_file(path, mimetype='.docx', attachment_filename=rename, as_attachment=True)
        except Exception as E:
            return "error"

    file_obj = request.files['file']
    if file_obj is None:
        return "nofile"

    print(file_obj.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj.filename)
    file_obj.save(file_path)

    document = Document()
    pdf=os.open('./upload/'+file_obj.filename,os.O_RDWR)
    print(pdf)
    fn = open(pdf,'rb')
    parser = PDFParser(fn)
    doc = PDFDocument()
    parser.set_document(doc)
    doc.set_parser(parser)
    resource = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(resource,laparams=laparams)
    interpreter = PDFPageInterpreter(resource,device)
    file_obj.filename = file_obj.filename.replace('.pdf', '')
    file_obj.filename = file_obj.filename + '.docx'
    for i in doc.get_pages():
        interpreter.process_page(i)
        layout = device.get_result()
        for out in layout:
            if hasattr(out,"get_text"):
                content = out.get_text().replace(u'\n', u'')
                document.add_paragraph(
                    content, style='ListBullet'
                )
            document.save(r'./download/result.docx')

    print ('处理完成')
    return "success"


@app.route('/wordCloud/<argv>/<ran>', methods=['post', 'get'])
def wordCloud(argv,ran):
    # pip install pdfminer3k
    # pip install python_docx
    if argv == "download":
        path = r'./download/result.png'
        rename = "result.png"
        try:
            return send_file(path, mimetype='.png', attachment_filename=rename, as_attachment=True)
        except Exception as E:
            return "error"

    file_obj = request.files['file']
    if file_obj is None:
        return "nofile"

    print(file_obj.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj.filename)
    file_obj.save(file_path)
    f = open('./upload/'+file_obj.filename,encoding = 'utf-8')
    text = f.read()
    print(text)
    f.close()
    wordlist = jieba.cut(text, cut_all=True)
    wl = " ".join(wordlist)
    print(wl)  # 输出分词之后的txt
    coloring = np.array(Image.open("./background.png"))

    # 设置停用词
    stopwords = set(STOPWORDS)
    stopwords.add("said")

    # 你可以通过 mask 参数 来设置词云形状
    # wc = WordCloud(background_color="white", max_words=2000, mask=coloring,
    #                 max_font_size=50, random_state=42,font_path='fangsong_GB2312.ttf')
    wc = WordCloud(background_color="white", max_words=2000, mask=coloring,
                   max_font_size=50, random_state=42, font_path='Hiragino Sans GB.ttf')
    wc.generate(wl)

    # create coloring from image
    image_colors = ImageColorGenerator(coloring)
    wc.to_file('./download/result.png')
    return "success"


@app.route('/ploydata/<argv>/<ran>', methods=['post', 'get'])
def polydata(argv,ran):
    # pip install pdfminer3k
    # pip install python_docx
    if argv == "download":
        path = r'./download/result.csv'
        rename = "result.csv"
        try:
            return send_file(path, mimetype='.csv', attachment_filename=rename, as_attachment=True)
        except Exception as E:
            return "error"

    file_obj = request.files['file']
    if file_obj is None:
        return "nofile"


    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj.filename)
    file_obj.save(file_path)


    csv_data = pd.read_csv(r'./upload/'+file_obj.filename, low_memory=False)  # 防止弹出警告
    data = pd.DataFrame(csv_data)
    for i in data.columns:
        for j in range(len(data)):
            if (np.isnan(data[i][j])):
                data[i][j] = ploy(data[i], j)

    data.to_csv('./download/result.csv', index=False, header=False)
    print('处理完成')
    return "success"


#决策树处理分类问题
@app.route('/decisionTree/<argv>/<ran>', methods=['post', 'get'])
def decisionTree(argv,ran):
    if argv == "download":
        path = r'./download/result.csv'
        rename = "result.csv"
        try:
            return send_file(path, mimetype='.csv', attachment_filename=rename, as_attachment=True)
        except Exception as E:
            return "error"
    file_obj = request.files.getlist("file")
    if len(file_obj) < 2:
        ret = "需要上传训练集和测试集以进行模型训练与测试"
        return ret
    train=[]
    test=[]
    print(file_obj)
    train_file = file_obj[0].filename
    test_file = file_obj[1].filename
    print(train_file)
    print(test_file)


    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj[0].filename)
    file_obj[0].save(file_path)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj[1].filename)
    file_obj[1].save(file_path)
    train = get_train_dataset('train.csv')
    test = get_test_dataset('test.csv')
    # print('test=',test)
    train_data, train_target, valid_data, valid_target = train_valid_data(train)
    model = tree.DecisionTreeClassifier(criterion='gini')
    print('train_data=',train_data)
    print('train_target=',train_target)
    model.fit(train_data,train_target)
    print("train score:", model.score(train_data, train_target))
    print("valid score:", model.score(valid_data, valid_target))
    predict = model.predict(test)
    print(predict)
    p_list = []
    for p in predict:
        list_tem = []
        list_tem.append(p)
        p_list.append(list_tem)
    print(p_list)
    with open(r'./download/result.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in p_list:
            # print(line)
            csv_writer.writerow(list(line))
    return "success"


@app.route('/randomForest/<argv>/<ran>', methods=['post', 'get'])
def randomForest(argv,ran):
    if argv == "download":
        path = r'./download/result.csv'
        rename = "result.csv"
        try:
            return send_file(path, mimetype='.csv', attachment_filename=rename, as_attachment=True)
        except Exception as E:
            return "error"
    file_obj = request.files.getlist("file")
    if len(file_obj) < 2:
        ret = "需要上传训练集和测试集以进行模型训练与测试"
        return ret
    train=[]
    test=[]
    print(file_obj)
    train_file = file_obj[0].filename
    test_file = file_obj[1].filename
    print(train_file)
    print(test_file)


    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj[0].filename)
    file_obj[0].save(file_path)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj[1].filename)
    file_obj[1].save(file_path)
    train = get_train_dataset('train.csv')
    test = get_test_dataset('test.csv')
    # print('test=',test)
    train_data, train_target, valid_data, valid_target = train_valid_data(train)
    model = RandomForestClassifier(n_estimators=100,
                                   bootstrap=True,
                                   max_features='sqrt')
    # Fit on training data
    model.fit(train_data, train_target)
    print('train_data=',train_data)
    print('train_target=',train_target)
    model.fit(train_data,train_target)
    print("train score:", model.score(train_data, train_target))
    print("valid score:", model.score(valid_data, valid_target))
    predict = model.predict(test)
    print(predict)
    p_list = []
    for p in predict:
        list_tem = []
        list_tem.append(p)
        p_list.append(list_tem)
    print(p_list)
    with open(r'./download/result.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in p_list:
            # print(line)
            csv_writer.writerow(list(line))
    return "success"

@app.route('/svm/<argv>/<ran>', methods=['post', 'get'])
def svm(argv,ran):
    if argv == "download":
        path = r'./download/result.csv'
        rename = "result.csv"
        try:
            return send_file(path, mimetype='.csv', attachment_filename=rename, as_attachment=True)
        except Exception as E:
            return "error"
    file_obj = request.files.getlist("file")
    if len(file_obj) < 2:
        ret = "需要上传训练集和测试集以进行模型训练与测试"
        return ret
    train=[]
    test=[]
    print(file_obj)
    train_file = file_obj[0].filename
    test_file = file_obj[1].filename
    print(train_file)
    print(test_file)


    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj[0].filename)
    file_obj[0].save(file_path)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj[1].filename)
    file_obj[1].save(file_path)
    train = get_train_dataset('train.csv')
    test = get_test_dataset('test.csv')
    # print('test=',test)
    train_data, train_target, valid_data, valid_target = train_valid_data(train)
    svclassifier = SVC(kernel='poly', degree=8)
    svclassifier.fit(train_data, train_target)
    print("train score:", svclassifier.score(train_data, train_target))
    print("valid score:", svclassifier.score(valid_data, valid_target))
    predict = svclassifier.predict(test)
    print(predict)
    p_list = []
    for p in predict:
        list_tem = []
        list_tem.append(p)
        p_list.append(list_tem)
    print(p_list)
    with open(r'./download/result.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in p_list:
            # print(line)
            csv_writer.writerow(list(line))
    return "success"


@app.route('/svr/<argv>/<ran>', methods=['post', 'get'])
def svr(argv,ran):
    if argv == "download":
        path = r'./download/result.csv'
        rename = "result.csv"
        try:
            return send_file(path, mimetype='.csv', attachment_filename=rename, as_attachment=True)
        except Exception as E:
            return "error"
    file_obj = request.files.getlist("file")
    if len(file_obj) < 2:
        ret = "需要上传训练集和测试集以进行模型训练与测试"
        return ret
    train=[]
    test=[]
    print(file_obj)
    train_file = file_obj[0].filename
    test_file = file_obj[1].filename
    print(train_file)
    print(test_file)


    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj[0].filename)
    file_obj[0].save(file_path)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj[1].filename)
    file_obj[1].save(file_path)
    train = get_train_dataset('train.csv')
    test = get_test_dataset('test.csv')
    # print('test=',test)
    train_data, train_target, valid_data, valid_target = train_valid_data(train)
    svrlassifier = SVR(kernel='rbf', C=20)
    # Fit on training data
    svrlassifier.fit(train_data, train_target)
    print("train score:", svrlassifier.score(train_data, train_target))
    print("valid score:", svrlassifier.score(valid_data, valid_target))
    predict = svrlassifier.predict(test)
    p_list = []
    for p in predict:
        list_tem = []
        list_tem.append(p)
        p_list.append(list_tem)
    print(p_list)
    with open(r'./download/result.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in p_list:
            # print(line)
            csv_writer.writerow(list(line))
    return "success"


@app.route('/linearRegression/<argv>/<ran>', methods=['post', 'get'])
def linearRegression(argv,ran):
    if argv == "download":
        path = r'./download/result.csv'
        rename = "result.csv"
        try:
            return send_file(path, mimetype='.csv', attachment_filename=rename, as_attachment=True)
        except Exception as E:
            return "error"
    file_obj = request.files.getlist("file")
    if len(file_obj) < 2:
        ret = "需要上传训练集和测试集以进行模型训练与测试"
        return ret
    train=[]
    test=[]
    print(file_obj)
    train_file = file_obj[0].filename
    test_file = file_obj[1].filename
    print(train_file)
    print(test_file)


    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj[0].filename)
    file_obj[0].save(file_path)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj[1].filename)
    file_obj[1].save(file_path)
    train = get_train_dataset('train.csv')
    test = get_test_dataset('test.csv')
    # print('test=',test)
    train_data, train_target, valid_data, valid_target = train_lr_valid_data(train)
    linreg = LinearRegression()
    model = linreg.fit(train_data, train_target)
    print(model)
    # 训练后模型截距
    print(linreg.intercept_)
    # 训练后模型权重（特征个数无变化）
    print(linreg.coef_)
    predict = linreg.predict(test)
    p_list = []
    for p in predict:
        list_tem = []
        list_tem.append(p)
        p_list.append(list_tem)
    print(p_list)
    with open(r'./download/result.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in p_list:
            # print(line)
            csv_writer.writerow(list(line))
    return "success"

@app.route('/logisticRegression/<argv>/<ran>', methods=['post', 'get'])
def logisticRegression(argv,ran):
    if argv == "download":
        path = r'./download/result.csv'
        rename = "result.csv"
        try:
            return send_file(path, mimetype='.csv', attachment_filename=rename, as_attachment=True)
        except Exception as E:
            return "error"
    file_obj = request.files.getlist("file")
    if len(file_obj) < 2:
        ret = "需要上传训练集和测试集以进行模型训练与测试"
        return ret
    train=[]
    test=[]
    print(file_obj)
    train_file = file_obj[0].filename
    test_file = file_obj[1].filename
    print(train_file)
    print(test_file)


    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj[0].filename)
    file_obj[0].save(file_path)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file_obj[1].filename)
    file_obj[1].save(file_path)
    train = get_train_dataset('train.csv')
    test = get_test_dataset('test.csv')
    # print('test=',test)
    train_data, train_target, valid_data, valid_target = train_valid_data(train)
    logreg = LogisticRegression()
    model = logreg.fit(train_data, train_target)
    print(model)
    predict = model.predict(test)
    print("train score:", model.score(train_data, train_target))
    print("valid score:", model.score(valid_data, valid_target))
    p_list = []
    for p in predict:
        list_tem = []
        list_tem.append(p)
        p_list.append(list_tem)
    print(p_list)
    with open(r'./download/result.csv', 'w', newline='') as new_file:
        csv_writer = csv.writer(new_file)
        for line in p_list:
            # print(line)
            csv_writer.writerow(list(line))
    return "success"

@app.route('/test', methods=['post', 'get'])
def test():
    return "success"

if __name__ == '__main__':
    app.run(host="0.0.0.0", debug=False, port=8025)
